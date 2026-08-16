"""Turn an uploaded career document into proposed KB chunks.

Two clearly separated steps:
1. extract_text  — deterministic: pull raw text out of pdf/docx/txt/md.
2. parse_to_chunks — the LLM's job: structure that text into accomplishment
   chunks. It is instructed to use ONLY what's in the document (no inventing);
   the user reviews and edits everything before it is saved.
"""
import io
import re
from typing import Callable

import pdfplumber
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.ai.llm import generate_json

# How much document text goes into one model call. A long CV is split across
# several calls rather than truncated: cutting at 12k characters silently threw
# away the second half of anyone's career. Kept small (not just "under the
# model's context window") because small local models silently under-extract
# well before they run out of room: llama3.2:3b given ~7,400 chars in one call
# returned 2 of the ~20 real chunks, with no error. Splitting into more,
# smaller calls is the mitigation, not a bigger context budget.
_SEGMENT_CHARS = 3000


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from a supported document."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return _from_pdf(data)
    if ext == "docx":
        return _from_docx(data)
    if ext in ("txt", "md"):
        return data.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: .{ext} (use pdf, docx, txt, or md)")


# Private-use-area codepoints: never real text, just icon glyphs (contact-info
# icons and the like) that some templates draw inline with the words around
# them. The model has no use for a phone-icon codepoint.
_PRIVATE_USE_AREA = re.compile(r"[\uE000-\uF8FF]")

# What a bullet marker looks like across the templates we've seen. Used both to
# rejoin a wrapped bullet's continuation line in `_join_wrapped_lines` and to
# find entry boundaries in `_entries`.
_BULLET = re.compile(r"^\s*[\u2022\u25AA\u25E6\-*]")

# pdfplumber's default word-gap threshold (3pt) is wider than the actual glyph
# gaps in some resume-template fonts (typically LaTeX output), which merges
# adjacent words into one run with no space between them at all. Narrowing it
# recovers real word boundaries; checked against a plain Helvetica PDF too, so
# this isn't just chasing one font.
_PDF_X_TOLERANCE = 0.5

# Some LaTeX resume templates render small-caps headings with a font whose
# ToUnicode table is simply wrong for the small lowercase-shaped glyphs, so
# extraction comes back with a single stray lowercase letter inside an
# otherwise all-caps word ("ENGiNEER"). A 3B-parameter model asked to fix this
# in the same breath as its other extraction rules left it uncorrected in
# testing, so it's handled deterministically here instead.
_SMALLCAPS_WORD = re.compile(r"[A-Za-z]{4,}")


def _fix_smallcaps_casing(text: str) -> str:
    def fix(match: re.Match) -> str:
        word = match.group(0)
        lower = [i for i, ch in enumerate(word) if ch.islower()]
        if len(lower) != 1:
            return word
        i = lower[0]
        return word[:i] + word[i].upper() + word[i + 1 :]

    return _SMALLCAPS_WORD.sub(fix, text)


def _join_wrapped_lines(page: "pdfplumber.page.Page") -> str:
    """Page text with each bullet's word-wrapped continuation rejoined onto it.

    A bullet that wraps to a second visual line comes back from `extract_text`
    as two plain lines with nothing marking the second as a continuation, which
    looks identical to a new heading line. Resumes reliably render that second
    line indented further right than the bullet marker itself, though (it lines
    up under the bullet's text, not the bullet), so `x0` — how far a line starts
    from the page's left edge — is what tells continuation and heading apart
    where the text alone can't.

    Assumes a single-column layout. A two-column resume's right column would
    read as consistently more indented than the left, so this would wrongly
    treat it as a run-on continuation — the same known multi-column limitation
    already accepted elsewhere in extraction, not a new gap.
    """
    lines = page.extract_text_lines(x_tolerance=_PDF_X_TOLERANCE)
    joined: list[str] = []
    bullet_indent: float | None = None
    for line in lines:
        text = line["text"]
        if _BULLET.match(text):
            joined.append(text)
            bullet_indent = line["x0"]
        elif bullet_indent is not None and line["x0"] > bullet_indent:
            joined[-1] = f"{joined[-1]} {text}"
        else:
            joined.append(text)
            bullet_indent = None
    return "\n".join(joined)


def _from_pdf(data: bytes) -> str:
    """Page text, plus any tables the page draws.

    `extract_text` alone reads a table as whatever order its cells happen to sit
    in, which for a two column skills grid is every other word. Pulling the
    tables out separately and writing each row on its own line at least keeps the
    row together, which is the unit that means something.
    """
    parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                page_text = _join_wrapped_lines(page)
                page_text = _PRIVATE_USE_AREA.sub("", page_text)
                parts.append(_fix_smallcaps_casing(page_text))
                for table in page.extract_tables() or []:
                    for row in table:
                        cells = [str(c).strip() for c in row if c and str(c).strip()]
                        if cells:
                            parts.append(" | ".join(cells))
    except Exception as e:
        raise ValueError(
            "Couldn't open that PDF, it might be locked or a little broken. "
            "Re-export it and give it another try."
        ) from e

    text = "\n".join(p for p in parts if p.strip())
    if not text.strip():
        # A PDF that opens fine but yields no text is almost always a scan or a
        # flattened image, not a bug in the extraction above — a different
        # message than the open-failure one because the fix is different too.
        raise ValueError(
            "I can open your PDF but I cannot read anything from it. Could it "
            "be an image based PDF or a text saved as a picture? Try exporting "
            "your PDF as a digital scannable PDF and we will take it from there."
        )
    return text


def _from_docx(data: bytes) -> str:
    """Everything in the document, in the order it is laid out.

    `Document.paragraphs` skips table cells entirely, and a resume that puts its
    skills, education or contact details in a table is common enough that this
    was the single biggest way text went missing: not truncated, never read.

    Walking the body's own children keeps paragraphs and tables interleaved in
    document order, so a heading still arrives immediately before the table it
    introduces.

    Headers and footers are appended at the end because that is where phone
    numbers and email addresses hide in a lot of templates.
    """
    doc = Document(io.BytesIO(data))
    lines = list(_iter_block_text(doc))

    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                if paragraph.text.strip():
                    lines.append(paragraph.text.strip())

    return "\n".join(lines)


def _iter_block_text(parent: DocxDocument):
    """Paragraph and table text from one container, in document order."""
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            text = Paragraph(child, parent).text.strip()
            if text:
                yield text
        elif child.tag == qn("w:tbl"):
            for row in Table(child, parent).rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                # Deduplicated because a merged cell reports its text once per
                # column it spans, which would otherwise repeat every entry.
                unique = list(dict.fromkeys(cells))
                if unique:
                    yield " | ".join(unique)


_SYSTEM = """You extract a candidate's career into structured accomplishment chunks for a resume knowledge base.

Rules:
- Use ONLY information present in the document. Never invent facts, metrics, employers, dates, or technologies.
- Create ONE chunk per distinct accomplishment or responsibility (split multi-part bullets).
- `title` is the person's role/job title/degree/project name (e.g. "Software Engineer"). `company` is the employer/institution/organization name (e.g. "Acme Corp"). Never put the same value in both, and never leave `title` blank.
- A single job or project usually produces several chunks, one per bullet. Repeat `title`, `company`, and `date_range` identically on EVERY chunk from that job or project, even though they don't change between bullets. Never state them on only the first chunk and omit them from the rest.
- `company` must be a name that literally appears in the text near that entry. Many personal/side projects have no employer at all — when none is named, `company` is null. Never guess a company from what the project is *about* (a project about cars is not automatically related to a car company); never pull `company` from inside the project's own title either.
- Choose `type` from: project, experience, leadership, achievement, skill, certification, education.
- Use `education` for degrees, schools and coursework, one chunk per qualification. Put the degree in `title`, the institution in `company` and the years in `date_range`.
- Fill `technologies` and `skills` only with items actually mentioned; otherwise use [].
- Fill `impact` only if a concrete outcome/metric is stated; otherwise null.
- `accomplishment` is REQUIRED on every chunk — it is the bullet's own text (what was done), and it is never empty or omitted. `context` is a different, optional field: a short note ABOUT the accomplishment (e.g. the project or initiative it was part of), not a substitute for it. A chunk with a `context` but no `accomplishment` is wrong.
- Keep `accomplishment` concise, truthful, and results-oriented.

Respond as JSON: {"chunks": [{"type","title","context","company","date_range","accomplishment","technologies":[],"skills":[],"impact"}]}"""


# The headings a resume actually uses, and what each one means in our own
# vocabulary.
#
# Ordered most specific first, and "experience" is deliberately last: it is the
# word most likely to turn up inside a combined heading, and a template that
# writes "Skills & Experience" means the skills section.
_SECTION_TYPES: tuple[tuple[str, str], ...] = (
    (r"educations?|academics?|qualifications?|schooling", "education"),
    (r"certifications?|certificates?|licen[cs]es?|courses?|coursework|training", "certification"),
    (r"skills?|technologies|technology|technical|tools|competenc(?:y|ies)", "skill"),
    (r"projects?|portfolio", "project"),
    (r"leadership|volunteering|volunteer|activities|responsibility|extracurricular", "leadership"),
    (r"achievements?|awards?|honou?rs?|accomplishments?|publications?", "achievement"),
    (r"experiences?|employment|internships?|professional|career", "experience"),
)

_SECTION_PATTERNS = tuple(
    (re.compile(rf"(?<![a-z]){words}(?![a-z])", re.IGNORECASE), kind)
    for words, kind in _SECTION_TYPES
)

# A heading is a label, not a sentence. Both bounds are what stop an ordinary
# line that happens to contain "education" from refiling everything after it.
_MAX_HEADING_CHARS = 40
_MAX_HEADING_WORDS = 4


def section_of(line: str) -> str | None:
    """The chunk type this line announces, if it is a section heading at all.

    Extraction flattens a document to plain text, which throws away the font size
    and the bold that made a heading look like one. All that survives is the
    shape of the line: short, standing alone, and naming a section. That shape is
    what this reads, and it is the only thing standing between "EDUCATION" and a
    degree being filed as another job.

    Shape first, keyword second, deliberately. Keyword alone would read "My
    education taught me to ship early" as a section break.
    """
    stripped = line.strip()
    if not stripped or len(stripped) > _MAX_HEADING_CHARS:
        return None
    if len(stripped.split()) > _MAX_HEADING_WORDS:
        return None
    for pattern, kind in _SECTION_PATTERNS:
        if pattern.search(stripped):
            return kind
    return None


def segment(text: str) -> list[str]:
    """Split a document into model-sized pieces, breaking at the cleanest seam.

    Where it breaks matters: cutting mid-bullet hands the model half an
    accomplishment and it faithfully records half an accomplishment. So it tries
    paragraph gaps first, then single line breaks, and only chops mid-line when a
    document offers no seam at all.

    Both fallbacks are load-bearing on real files. A Windows text file separates
    paragraphs with CRLF, and PDF text extraction routinely comes back as one
    long run of single newlines — either would otherwise sail past a
    blank-line-only split and reach the model whole.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if len(text) <= _SEGMENT_CHARS:
        return [text] if text else []

    return _pack(_pieces(text))


def segments_with_sections(text: str) -> list[tuple[str, str | None]]:
    """Each segment, paired with the section heading that was open when it began.

    A long resume is split into several model calls, and the split lands wherever
    the text allows. A segment starting halfway through the education list used
    to arrive with no clue what it was, so the model guessed, and it guessed the
    most common thing in a resume, which is jobs. Carrying the heading forward is
    what stops the second half of a section being filed as something else.
    """
    pieces = segment(text)
    out: list[tuple[str, str | None]] = []
    current: str | None = None
    for piece in pieces:
        out.append((piece, current))
        for line in piece.split("\n"):
            found = section_of(line)
            if found:
                current = found
    return out



# A one-line summary some project entries trail their bullets with. Shape-wise
# it's a non-bullet line right after a run of bullets — identical to a new
# heading — but it belongs to the entry ending, not the one about to start.
_TRAILING_SUMMARY = re.compile(r"^\s*Tech(?:nical)? Stack\s*:", re.IGNORECASE)


def _entries(text: str) -> list[str]:
    """Group each heading with the bullets under it, so the two never separate.

    PDF-extracted resumes have no blank line between one job and the next — a
    bullet is just the next line after the last one. Packing by raw line would
    happily split a job's heading (company, title, dates) into one segment and
    its bullets into the next, and the model receiving only the bullets has no
    company to attach them to. A new entry starts wherever a non-bullet line
    follows a bullet: that shape is what separates one job/project from the
    next, since the heading lines before the first bullet belong to it.

    A trailing "Tech Stack:" summary line is the one common exception to that
    shape: it comes right after the bullets too, but it describes the entry
    that's ending, not the one about to start. Treating it like a bullet for
    grouping purposes (append, don't flush) keeps it attached to its own entry
    and defers the flush to the heading that actually follows it.
    """
    entries: list[str] = []
    current: list[str] = []
    entry_open = False
    for line in text.split("\n"):
        stays_open = bool(_BULLET.match(line)) or bool(_TRAILING_SUMMARY.match(line))
        if not stays_open and entry_open:
            entries.append("\n".join(current))
            current = []
            entry_open = False
        current.append(line)
        entry_open = entry_open or stays_open
    if current:
        entries.append("\n".join(current))
    return entries


def _pieces(text: str) -> list[str]:
    """Break text into units no larger than one segment, cleanest seam first."""
    units = text.split("\n\n")
    if all(len(u) <= _SEGMENT_CHARS for u in units):
        return units

    units = _entries(text)
    if all(len(u) <= _SEGMENT_CHARS for u in units):
        return units

    units = [line for unit in units for line in unit.split("\n")]
    if all(len(u) <= _SEGMENT_CHARS for u in units):
        return units

    return [
        unit[i : i + _SEGMENT_CHARS]
        for unit in units
        for i in range(0, len(unit), _SEGMENT_CHARS)
    ]


def _pack(units: list[str]) -> list[str]:
    """Greedily fill segments so a long document is as few model calls as possible."""
    segments: list[str] = []
    current = ""
    for unit in units:
        if not unit.strip():
            continue
        if current and len(current) + len(unit) + 2 > _SEGMENT_CHARS:
            segments.append(current)
            current = unit
        else:
            current = f"{current}\n\n{unit}" if current else unit
    if current:
        segments.append(current)
    return segments


def parse_to_chunks(
    text: str,
    kind_hint: str | None = None,
    on_segment: Callable[[], None] | None = None,
) -> list[dict]:
    """Structure document text into chunk dicts, one model call per segment.

    No timeout: a 30-page CV on a laptop model is slow, not broken, and the
    caller runs this in the background with a progress bar for exactly that
    reason. `on_segment` fires after each call so that bar can move.
    """
    hint = f"\n\nThis document is the candidate's: {kind_hint}." if kind_hint else ""
    chunks: list[dict] = []
    for piece, section in segments_with_sections(text):
        # Told, not left to be inferred: a piece that begins mid-section has no
        # heading in it to read.
        carried = (
            f"\n\nThis continues the {section} section of the document."
            if section
            else ""
        )
        # A resume can yield many chunks, so this needs far more room than extraction.
        result = generate_json(
            _SYSTEM, f"Document:{hint}{carried}\n\n{piece}", timeout=None, max_tokens=4000
        )
        raw = result.get("chunks", []) if isinstance(result, dict) else []
        chunks.extend(
            _normalize(c, section)
            for c in raw
            if isinstance(c, dict) and (c.get("accomplishment") or c.get("context"))
        )
        if on_segment:
            on_segment()
    return chunks


_VALID_TYPES = {
    "project", "experience", "leadership", "achievement", "skill",
    "certification", "education",
}

# LLMs often emit the *string* "null"/"none"/"n/a" instead of a real null.
_NULLISH = {"", "null", "none", "n/a", "na", "unknown", "not specified"}


def _opt(value, limit: int) -> str | None:
    """Optional text field: strip, drop nullish placeholders, cap length."""
    if value is None:
        return None
    text = str(value).strip()
    return text[:limit] if text.lower() not in _NULLISH else None


def _normalize(c: dict, section: str | None = None) -> dict:
    """Coerce a raw LLM chunk into our exact shape with safe defaults.

    When the model returns a type we do not recognise, the section the text came
    from decides, and only a document with no heading at all falls back to
    "experience". That fallback used to be unconditional, and since `type` is the
    single thing that picks the heading a chunk is printed under, it was the
    mechanism by which a degree ended up in the middle of someone's job history.
    """
    ctype = str(c.get("type", "")).lower().strip()
    if ctype not in _VALID_TYPES:
        ctype = section if section in _VALID_TYPES else "experience"
    # The model sometimes puts the descriptive text in `context` and leaves
    # `accomplishment` off the chunk entirely rather than filling both. When
    # that happens, `context` *is* the accomplishment - treating it as a
    # separate, still-empty field would silently drop the whole chunk instead.
    accomplishment = _opt(c.get("accomplishment"), 4000)
    context = _opt(c.get("context"), 300)
    if not accomplishment:
        accomplishment = _opt(c.get("context"), 4000)
        context = None
    return {
        "type": ctype,
        "title": _opt(c.get("title"), 300) or "Untitled",
        "context": context,
        "company": _opt(c.get("company"), 200),
        "date_range": _opt(c.get("date_range"), 100),
        "accomplishment": accomplishment or "",
        "technologies": [t for t in (_opt(x, 100) for x in (c.get("technologies") or [])) if t],
        "skills": [s for s in (_opt(x, 100) for x in (c.get("skills") or [])) if s],
        "impact": _opt(c.get("impact"), 500),
    }
