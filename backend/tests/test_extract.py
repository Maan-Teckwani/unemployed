"""What actually comes out of an uploaded document.

`segment()` was the only extraction thing under test, and it only checks sizes.
Nothing tested `extract_text` for either real format, which is how a DOCX losing
every table went unnoticed: python-docx's `Document.paragraphs` skips table cells
entirely, so a resume that laid its skills out in a table had them deleted before
the model ever saw the file. Not truncated. Never read.
"""
import io

import pytest
from docx import Document
from fpdf import FPDF

from app.ai.parse import extract_text


def docx_bytes(build) -> bytes:
    doc = Document()
    build(doc)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---- DOCX -------------------------------------------------------------------
def test_a_table_is_read(tmp_path) -> None:
    def build(doc):
        doc.add_paragraph("SKILLS")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Languages"
        table.cell(0, 1).text = "Python, Go"
        table.cell(1, 0).text = "Databases"
        table.cell(1, 1).text = "Postgres, SQLite"

    text = extract_text("cv.docx", docx_bytes(build))
    for expected in ("Languages", "Python, Go", "Databases", "Postgres, SQLite"):
        assert expected in text, f"{expected!r} was dropped from the document"


def test_a_table_keeps_its_rows_together(tmp_path) -> None:
    # A row is the unit that means something: the label and the value belong on
    # one line or the pairing is lost.
    def build(doc):
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "B.Tech Computer Science"
        table.cell(0, 1).text = "2021 to 2025"

    text = extract_text("cv.docx", docx_bytes(build))
    line = next(line for line in text.splitlines() if "B.Tech" in line)
    assert "2021 to 2025" in line


def test_paragraphs_and_tables_stay_in_document_order() -> None:
    # A heading has to arrive before the table it introduces, or the section it
    # names is attached to the wrong content.
    def build(doc):
        doc.add_paragraph("EXPERIENCE")
        t1 = doc.add_table(rows=1, cols=1)
        t1.cell(0, 0).text = "Worked at Acme"
        doc.add_paragraph("EDUCATION")
        t2 = doc.add_table(rows=1, cols=1)
        t2.cell(0, 0).text = "Studied at Some University"

    text = extract_text("cv.docx", docx_bytes(build))
    assert (
        text.index("EXPERIENCE")
        < text.index("Worked at Acme")
        < text.index("EDUCATION")
        < text.index("Studied at Some University")
    )


def test_header_and_footer_text_is_kept() -> None:
    # Templates routinely put the phone number and email in the header, which
    # `Document.paragraphs` never sees.
    def build(doc):
        doc.add_paragraph("Body text")
        doc.sections[0].header.paragraphs[0].text = "me@example.com | +91 90000 00000"

    text = extract_text("cv.docx", docx_bytes(build))
    assert "me@example.com" in text


def test_a_merged_cell_is_not_repeated() -> None:
    # A merged cell reports its text once per column it spans, which would
    # otherwise print the same entry several times on one line.
    def build(doc):
        table = doc.add_table(rows=1, cols=3)
        table.cell(0, 0).merge(table.cell(0, 2)).text = "One wide heading"

    text = extract_text("cv.docx", docx_bytes(build))
    assert text.count("One wide heading") == 1


def test_ordinary_paragraphs_still_work() -> None:
    def build(doc):
        doc.add_paragraph("First line")
        doc.add_paragraph("Second line")

    text = extract_text("cv.docx", docx_bytes(build))
    assert "First line" in text and "Second line" in text


# ---- PDF --------------------------------------------------------------------
def pdf_bytes(rows: list[tuple[str, str]], intro: str = "") -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "", 11)
    if intro:
        pdf.cell(0, 8, intro, new_x="LMARGIN", new_y="NEXT")
    # Bordered cells, because pdfplumber finds tables by their ruling lines.
    for left, right in rows:
        pdf.cell(60, 8, left, border=1)
        pdf.cell(90, 8, right, border=1, new_x="LMARGIN", new_y="NEXT")
    return bytes(pdf.output())


def test_pdf_body_text_is_read() -> None:
    text = extract_text("cv.pdf", pdf_bytes([], intro="Plain line of the resume"))
    assert "Plain line of the resume" in text


def test_pdf_table_cells_are_read() -> None:
    data = [("Languages", "Python, Go"), ("Databases", "Postgres")]
    text = extract_text("cv.pdf", pdf_bytes(data, intro="SKILLS"))
    assert "SKILLS" in text
    for cell in ("Languages", "Python, Go", "Databases", "Postgres"):
        assert cell in text, f"{cell!r} was dropped from the PDF"


def test_pdf_table_rows_are_recoverable() -> None:
    # extract_text alone reads a table in whatever order the cells sit in, which
    # for a two column grid is every other word. Pulling the tables out
    # separately is what keeps a row readable as a row.
    text = extract_text("cv.pdf", pdf_bytes([("Languages", "Python, Go")]))
    assert any(
        "Languages" in line and "Python, Go" in line for line in text.splitlines()
    ), "no single line holds the whole row"


# ---- Everything else --------------------------------------------------------
def test_plain_text_is_passed_through() -> None:
    assert "hello" in extract_text("notes.txt", b"hello")
    assert "hello" in extract_text("notes.md", b"# hello")


def test_an_unsupported_type_says_so() -> None:
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text("resume.pages", b"whatever")
